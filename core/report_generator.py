
import pandas as pd
import os
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches
import markdown
from typing import Dict, Any, Optional

def get_chinese_font():
    font_paths = [
        'C:/Windows/Fonts/simsun.ttc',
        'C:/Windows/Fonts/msyh.ttc',
        '/Library/Fonts/SimHei.ttf',
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
    ]
    for path in font_paths:
        if os.path.exists(path):
            return path
    return None

class ReportGenerator:
    @staticmethod
    def generate_markdown_report(data: Dict[str, Any], output_path: str) -> bool:
        try:
            md_content = f"""# 数据分析报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 一、数据概览

- **数据行数**: {data.get('shape', [0, 0])[0]}
- **数据列数**: {data.get('shape', [0, 0])[1]}

## 二、数据列信息

| 列名 | 数据类型 | 缺失值数量 |
|------|----------|------------|
"""
            columns = data.get('columns', [])
            dtypes = data.get('dtypes', {})
            missing_count = data.get('missing_count', {})
            
            for col in columns:
                md_content += f"| {col} | {dtypes.get(col, 'unknown')} | {missing_count.get(col, 0)} |\n"
            
            md_content += """

## 三、描述性统计

"""
            descriptive = data.get('descriptive', pd.DataFrame())
            if not descriptive.empty:
                md_content += descriptive.to_markdown()
            
            md_content += """

## 四、相关性分析

"""
            correlation = data.get('correlation', pd.DataFrame())
            if not correlation.empty:
                md_content += correlation.to_markdown()
            
            md_content += f"""

## 五、分析说明

- **重复行数**: {data.get('duplicate_count', 0)}
- **分析状态**: 完成

---

*本报告由 DataAnalyser 自动生成*
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            return True
        except Exception as e:
            return False

    @staticmethod
    def generate_word_report(data: Dict[str, Any], output_path: str) -> bool:
        try:
            doc = Document()
            
            title_style = doc.styles['Heading 1']
            title_style.font.size = Pt(16)
            title_style.bold = True
            
            heading_style = doc.styles['Heading 2']
            heading_style.font.size = Pt(14)
            
            doc.add_heading('数据分析报告', 0)
            doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')
            
            doc.add_heading('一、数据概览', level=1)
            doc.add_paragraph(f'数据行数: {data.get("shape", [0, 0])[0]}')
            doc.add_paragraph(f'数据列数: {data.get("shape", [0, 0])[1]}')
            
            doc.add_heading('二、数据列信息', level=1)
            columns = data.get('columns', [])
            dtypes = data.get('dtypes', {})
            missing_count = data.get('missing_count', {})
            
            table_data = [['列名', '数据类型', '缺失值数量']]
            for col in columns:
                table_data.append([col, str(dtypes.get(col, 'unknown')), str(missing_count.get(col, 0))])
            
            table = doc.add_table(rows=len(table_data), cols=3)
            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = cell
            
            doc.add_heading('三、描述性统计', level=1)
            descriptive = data.get('descriptive', pd.DataFrame())
            if not descriptive.empty:
                desc_table = doc.add_table(rows=len(descriptive)+1, cols=len(descriptive.columns)+1)
                desc_table.cell(0, 0).text = '特征'
                for j, col in enumerate(descriptive.columns):
                    desc_table.cell(0, j+1).text = str(col)
                for i, (idx, row) in enumerate(descriptive.iterrows()):
                    desc_table.cell(i+1, 0).text = str(idx)
                    for j, val in enumerate(row):
                        desc_table.cell(i+1, j+1).text = str(round(val, 2) if isinstance(val, float) else val)
            
            doc.add_heading('四、相关性分析', level=1)
            correlation = data.get('correlation', pd.DataFrame())
            if not correlation.empty:
                corr_table = doc.add_table(rows=len(correlation)+1, cols=len(correlation.columns)+1)
                corr_table.cell(0, 0).text = ''
                for j, col in enumerate(correlation.columns):
                    corr_table.cell(0, j+1).text = str(col)
                for i, (idx, row) in enumerate(correlation.iterrows()):
                    corr_table.cell(i+1, 0).text = str(idx)
                    for j, val in enumerate(row):
                        corr_table.cell(i+1, j+1).text = str(round(val, 4))
            
            doc.add_heading('五、分析说明', level=1)
            doc.add_paragraph(f'重复行数: {data.get("duplicate_count", 0)}')
            doc.add_paragraph('分析状态: 完成')
            
            doc.add_paragraph('')
            doc.add_paragraph('---')
            doc.add_paragraph('本报告由 DataAnalyser 自动生成')
            
            doc.save(output_path)
            return True
        except Exception as e:
            return False

    @staticmethod
    def generate_pdf_report(data: Dict[str, Any], output_path: str) -> bool:
        try:
            chinese_font_path = get_chinese_font()
            if chinese_font_path:
                if 'simsun' in chinese_font_path.lower():
                    pdfmetrics.registerFont(TTFont('SimSun', chinese_font_path, subfontIndex=0))
                else:
                    pdfmetrics.registerFont(TTFont('SimSun', chinese_font_path))
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []
            
            title_style = ParagraphStyle(
                'Title', 
                parent=styles['Heading1'], 
                fontSize=18, 
                alignment=1,
                fontName='SimSun' if chinese_font_path else 'Helvetica-Bold'
            )
            heading_style = ParagraphStyle(
                'Heading', 
                parent=styles['Heading2'], 
                fontSize=14,
                fontName='SimSun' if chinese_font_path else 'Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontName='SimSun' if chinese_font_path else 'Helvetica'
            )
            italic_style = ParagraphStyle(
                'ChineseItalic',
                parent=styles['Italic'],
                fontName='SimSun' if chinese_font_path else 'Helvetica-Oblique'
            )
            
            elements.append(Paragraph('数据分析报告', title_style))
            elements.append(Paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', normal_style))
            elements.append(Spacer(1, 0.2*inch))
            
            elements.append(Paragraph('一、数据概览', heading_style))
            elements.append(Paragraph(f'数据行数: {data.get("shape", [0, 0])[0]}', normal_style))
            elements.append(Paragraph(f'数据列数: {data.get("shape", [0, 0])[1]}', normal_style))
            elements.append(Spacer(1, 0.1*inch))
            
            elements.append(Paragraph('二、数据列信息', heading_style))
            columns = data.get('columns', [])
            dtypes = data.get('dtypes', {})
            missing_count = data.get('missing_count', {})
            
            table_data = [['列名', '数据类型', '缺失值数量']]
            for col in columns:
                table_data.append([col, str(dtypes.get(col, 'unknown')), str(missing_count.get(col, 0))])
            
            table_font = 'SimSun' if chinese_font_path else 'Helvetica'
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,-1), table_font),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                ('GRID', (0,0), (-1,-1), 1, colors.black)
            ]))
            elements.append(table)
            elements.append(Spacer(1, 0.1*inch))
            
            elements.append(Paragraph('三、描述性统计', heading_style))
            descriptive = data.get('descriptive', pd.DataFrame())
            if not descriptive.empty:
                desc_data = [['特征'] + descriptive.columns.tolist()]
                for idx, row in descriptive.iterrows():
                    row_data = [str(idx)] + [str(round(val, 2) if isinstance(val, float) else val) for val in row]
                    desc_data.append(row_data)
                
                desc_table = Table(desc_data)
                desc_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,-1), table_font),
                    ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                    ('GRID', (0,0), (-1,-1), 1, colors.black)
                ]))
                elements.append(desc_table)
            
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph('四、相关性分析', heading_style))
            correlation = data.get('correlation', pd.DataFrame())
            if not correlation.empty:
                corr_data = [[''] + correlation.columns.tolist()]
                for idx, row in correlation.iterrows():
                    row_data = [str(idx)] + [str(round(val, 4)) for val in row]
                    corr_data.append(row_data)
                
                corr_table = Table(corr_data)
                corr_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('BACKGROUND', (0,1), (0,-1), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('TEXTCOLOR', (0,0), (0,-1), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,-1), table_font),
                    ('BACKGROUND', (1,1), (-1,-1), colors.beige),
                    ('GRID', (0,0), (-1,-1), 1, colors.black)
                ]))
                elements.append(corr_table)
            
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph('五、分析说明', heading_style))
            elements.append(Paragraph(f'重复行数: {data.get("duplicate_count", 0)}', normal_style))
            elements.append(Paragraph('分析状态: 完成', normal_style))
            
            elements.append(Spacer(1, 0.2*inch))
            elements.append(Paragraph('---', normal_style))
            elements.append(Paragraph('本报告由 DataAnalyser 自动生成', italic_style))
            
            doc.build(elements)
            return True
        except Exception as e:
            return False

    @staticmethod
    def generate_report(data: Dict[str, Any], output_path: str, format_type: str = 'pdf') -> bool:
        if format_type == 'pdf':
            return ReportGenerator.generate_pdf_report(data, output_path)
        elif format_type == 'word':
            return ReportGenerator.generate_word_report(data, output_path)
        elif format_type == 'markdown':
            return ReportGenerator.generate_markdown_report(data, output_path)
        return False
